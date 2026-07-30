from pathlib import Path
import re
import time
import csv
import urllib.parse
import urllib.request
import xml.etree.ElementTree as ET

from docx import Document


# ============================================================
# 1. Paths
# ============================================================
base_dir = Path(r"C:\Users\a2618\Desktop\Model_Paper")

input_docx = base_dir / "Disease_Associations_PaperLinks_Only.docx"
output_docx = base_dir / "Disease_Associations_With_Paper_Titles.docx"
cache_file = base_dir / "pmid_title_cache.csv"

# Put your email here for NCBI Entrez
your_email = "your_email@example.com"

# Batch size for PubMed requests
batch_size = 100

# Sleep between requests to avoid hitting NCBI too quickly
sleep_seconds = 0.4


# ============================================================
# 2. Helper functions
# ============================================================
def extract_pmid(text):
    """
    Extract PMID from a text line like:
    References: PMID:15841208
    """
    m = re.search(r"PMID\s*:\s*(\d+)", text, flags=re.IGNORECASE)
    if m:
        return m.group(1)
    return None


def clean_title(title):
    """
    Clean PubMed title text.
    """
    if title is None:
        return ""
    title = re.sub(r"\s+", " ", title).strip()
    return title


def load_cache(cache_path):
    """
    Load existing PMID-title cache if it exists.
    """
    cache = {}

    if cache_path.exists():
        with open(cache_path, "r", newline="", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                pmid = str(row.get("PMID", "")).strip()
                title = str(row.get("Title", "")).strip()
                if pmid and title:
                    cache[pmid] = title

    return cache


def save_cache(cache_path, cache):
    """
    Save PMID-title cache.
    """
    with open(cache_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["PMID", "Title"])
        writer.writeheader()
        for pmid in sorted(cache.keys(), key=lambda x: int(x)):
            writer.writerow({
                "PMID": pmid,
                "Title": cache[pmid]
            })


def fetch_pubmed_titles(pmids, email, batch_size=100, sleep_seconds=0.4):
    """
    Fetch paper titles from PubMed using NCBI EFetch.
    Returns dictionary:
        PMID -> Title
    """
    titles = {}

    pmids = list(dict.fromkeys([str(x).strip() for x in pmids if str(x).strip()]))

    for i in range(0, len(pmids), batch_size):
        batch = pmids[i:i + batch_size]

        params = {
            "db": "pubmed",
            "id": ",".join(batch),
            "retmode": "xml",
            "email": email,
            "tool": "replace_disease_with_pubmed_title_script"
        }

        url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?" + urllib.parse.urlencode(params)

        print(f"Fetching PubMed titles {i + 1} to {i + len(batch)} of {len(pmids)}...")

        try:
            with urllib.request.urlopen(url, timeout=60) as response:
                xml_data = response.read()

            root = ET.fromstring(xml_data)

            for article in root.findall(".//PubmedArticle"):
                pmid_node = article.find(".//MedlineCitation/PMID")
                title_node = article.find(".//Article/ArticleTitle")

                if pmid_node is None:
                    continue

                pmid = pmid_node.text.strip()

                if title_node is not None:
                    # Handles titles with nested tags
                    title_text = "".join(title_node.itertext())
                    title_text = clean_title(title_text)
                else:
                    title_text = ""

                if title_text:
                    titles[pmid] = title_text

        except Exception as e:
            print(f"WARNING: Failed to fetch batch starting at PMID {batch[0]}")
            print(e)

        time.sleep(sleep_seconds)

    return titles


def replace_paragraph_with_title(paragraph, title):
    """
    Replace paragraph content with:
    Title: paper title

    Keeps the paragraph itself, style, indentation, spacing, etc.
    """
    # Remove existing runs
    for run in paragraph.runs:
        run.text = ""

    # Add new content
    run1 = paragraph.add_run("Title:")
    run1.bold = True

    run2 = paragraph.add_run(f" {title}")


# ============================================================
# 3. Read DOCX and collect all PMIDs
# ============================================================
doc = Document(input_docx)

paragraphs = list(doc.paragraphs)

all_pmids = []

for p in paragraphs:
    pmid = extract_pmid(p.text)
    if pmid:
        all_pmids.append(pmid)

all_pmids = list(dict.fromkeys(all_pmids))

print(f"Total unique PMIDs found: {len(all_pmids)}")


# ============================================================
# 4. Fetch titles from PubMed, using cache if available
# ============================================================
title_cache = load_cache(cache_file)

missing_pmids = [pmid for pmid in all_pmids if pmid not in title_cache]

print(f"PMIDs already in cache: {len(title_cache)}")
print(f"PMIDs needing PubMed lookup: {len(missing_pmids)}")

if missing_pmids:
    new_titles = fetch_pubmed_titles(
        missing_pmids,
        email=your_email,
        batch_size=batch_size,
        sleep_seconds=sleep_seconds
    )

    title_cache.update(new_titles)
    save_cache(cache_file, title_cache)

print(f"Total PMID titles available: {len(title_cache)}")


# ============================================================
# 5. Replace Disease lines with paper titles
# ============================================================
replace_count = 0
missing_title_count = 0

for i, p in enumerate(paragraphs):
    text = p.text.strip()

    # Only replace Disease lines
    if not text.startswith("Disease:"):
        continue

    # Find the next References line after this Disease line
    pmid = None

    for j in range(i + 1, min(i + 6, len(paragraphs))):
        next_text = paragraphs[j].text.strip()

        if next_text.startswith("References:"):
            pmid = extract_pmid(next_text)
            break

        # Stop if another Disease/Species/Gene block starts before a Reference line
        if next_text.startswith("Disease:") or next_text.startswith("Species "):
            break

    if pmid is None:
        continue

    title = title_cache.get(pmid, "").strip()

    if title:
        replace_paragraph_with_title(p, title)
        replace_count += 1
    else:
        # Keep original Disease line if title cannot be found
        missing_title_count += 1
        print(f"WARNING: No title found for PMID:{pmid}")


# ============================================================
# 6. Save output DOCX
# ============================================================
doc.save(output_docx)

print("Done.")
print(f"Replaced Disease lines: {replace_count}")
print(f"Missing titles: {missing_title_count}")
print(f"Output saved to: {output_docx}")
print(f"Cache saved to: {cache_file}")
